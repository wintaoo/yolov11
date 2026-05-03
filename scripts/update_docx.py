"""Generate updated CV检测系统2.0.docx with current system state."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('海南机器管招投标项目\n建筑图纸分析系统 v2.0')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('项目使用说明书  ·  2026年5月更新')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

# ============================================================
# 一、系统设计文档
# ============================================================
add_heading('一、系统设计文档', 1)

add_heading('1.1 系统概述', 2)
add_para('建筑图纸分析系统是面向施工招投标场景的智能分析平台，服务于海南机器管招投标项目。系统基于计算机视觉与大模型技术，提供两条核心处理管线：')
add_para('图片检测管线：单张图片 -> YOLO目标检测 -> 施工规范检查 -> 可视化标注', size=10)
add_para('文档分析管线：.docx文档 -> 图片提取 -> 多模态AI分析 -> 分类+评估 -> 汇总报告', size=10)

add_heading('1.2 系统架构', 2)
arch_text = """
浏览器前端 (Vue 3 + Element Plus)
  DetectionPanel | DocumentAnalysis | App.vue
  SSE实时进度 | Markdown渲染 | 队列状态可视化
       |  HTTP REST + SSE
Flask 后端 (:5000)
  detection.py          docx.py (Blueprint)
  /api/detect           /api/docx/*  (12 endpoints)
       |                      |
  Services Layer
  detection.py    rules_checker.py    docx_service.py
  (YOLOv11)       (Shapely 12条规则)   (OPC解析+图片压缩)
  ai_analysis_service.py
  (GLM-4.6V主 + GLM-4.5V备, 5 Key轮换)
  (图片MD5缓存, 压缩图优先, 原图>5MB跳过, 75s超时, 2次重试)
       |
  Data Layer: tasks/{id}/*.json | cache_ai/ | models/
       |
  SiliconFlow API
  GLM-4.6V (视觉主) | GLM-4.5V (视觉备) | Qwen2.5-7B (文本汇总)
"""
for line in arch_text.strip().split('\n'):
    add_para(line)

add_heading('1.3 AI分析服务详解', 2)

add_para('请求策略：', bold=True)
add_para('1. 图片压缩：长边 max 2048px, JPEG Q=85 -> base64')
add_para('2. 候选列表：[压缩图] + [原图（仅<=5MB，超大直接跳过）]')
add_para('3. Key + Model 随机选择 -> SiliconFlow Chat Completions API')
add_para('4. 每候选最多2次尝试（MAX_RETRIES_PER_CANDIDATE=1），超时75s')
add_para('5. 失败降级：HTTP 403->拉黑Key+Model；Timeout->切换候选；均失败->关键词匹配')

add_para('')
add_para('缓存机制：', bold=True)
add_para('图片MD5作为key，结果JSON序列化到 backend/cache_ai/。同图跨文档复用，仅缓存成功结果。')

add_para('')
add_para('JSON解析容错：', bold=True)
add_para('括号匹配提取 -> 多轮修复（去尾逗号）-> 关键词匹配兜底。evaluation字段通过_safe_str()统一处理dict/string混排。')

add_para('')
add_para('汇总报告策略（两级）：', bold=True)
add_para('1. 数据驱动报告：代码计算，100%准确（类型统计、评级分布、关键参数汇总）')
add_para('2. AI补充评审：调用文本模型生成简短意见，失败不影响主报告完整性')

add_heading('1.4 模型配置', 2)
add_table(
    ['配置变量', '用途', '当前值'],
    [
        ['SILICONFLOW_VISION_MODEL', '主视觉模型', 'zai-org/GLM-4.6V'],
        ['SILICONFLOW_VISION_MODELS', '备选视觉模型', 'zai-org/GLM-4.5V'],
        ['SILICONFLOW_SUMMARY_MODEL', '汇总文本模型', 'Qwen/Qwen2.5-7B-Instruct'],
        ['SILICONFLOW_API_KEYS', 'API密钥（逗号分隔）', '5个Key轮换'],
        ['IMAGE_COMPRESSION_ENABLED', '图片压缩', 'true'],
        ['IMAGE_MAX_DIMENSION', '压缩长边', '2048px'],
        ['AI_CACHE_ENABLED', 'AI缓存', 'true'],
    ]
)

add_heading('1.5 支持的图纸类型（13种）', 2)
add_table(
    ['序号', '类型', '典型内容'],
    [
        ['1', '总平面布置图', '施工区、办公区、生活区、道路、大门等整体布局'],
        ['2', '基础结构图', '桩基、筏板、承台等基础构件'],
        ['3', '主体结构图', '框架、剪力墙、柱网、塔吊、施工电梯等'],
        ['4', '土方工程图', '基坑开挖、支护、出土方向等'],
        ['5', '进度计划图', '横道图、网络图、甘特图、工期安排'],
        ['6', '施工计划图', '施工部署、施工组织、施工段划分'],
        ['7', '施工分区图', '流水段、施工区划分'],
        ['8', '分区规划图', '分区作业、流水施工'],
        ['9', '临时用电布置图', '配电房、配电箱、监控点位'],
        ['10', '临时用水布置图', '给水、排水、用水设施'],
        ['11', '临建设施平面布置图', '办公区、宿舍、食堂、门卫等'],
        ['12', '装饰装修图', '外立面、内装、精装修'],
        ['13', '周边环境图', '绿化、景观、周边道路衔接'],
    ]
)

add_heading('1.6 数据持久化', 2)
add_para('任务数据以JSON文件存储，目录结构：')
add_para('backend/tasks/{task_id}/')
add_para('  task.json          # 元数据（文件名、状态、图片列表）')
add_para('  results.json       # 分析结果 {编号: 结果对象}')
add_para('  summary.md         # Markdown汇总报告')
add_para('  analysis_report.md # 可读报告')
add_para('  images/            # 提取的图片文件')

add_heading('1.7 前端组件树', 2)
add_para('App.vue')
add_para('  AI免责声明条（橘色警示）')
add_para('  Tab切换（图片检测 / 文档分析）')
add_para('  DetectionPanel.vue')
add_para('    工具栏（文件夹选择、导航、清除）')
add_para('    对比视图（原图 vs 检测结果）')
add_para('    胶片导航 + 侧栏（检测/统计/列表/规范检查/导出）')
add_para('  DocumentAnalysis.vue')
add_para('    上传区（历史任务列表 + 上传卡片）')
add_para('    状态栏（任务ID/图片数/已分析/待分析）')
add_para('    进度条 + 异常横幅')
add_para('    图片网格（队列状态点：灰等待/紫分析中/绿完成/红异常）')
add_para('    详情面板（预览/分析按钮/AI结果）')
add_para('    汇总报告（Markdown渲染 + 导出PDF按钮）')
add_para('    分类统计')

# ============================================================
# 二、使用说明书
# ============================================================
add_heading('二、使用说明书', 1)

add_heading('2.1 启动与停止', 2)
add_para('启动：双击 start.bat（自动激活venv -> 后端:5000 -> 前端:8080）')
add_para('停止：双击 stop.bat 或关闭命令行窗口')
add_para('访问：http://localhost:8080')

add_heading('2.2 图片检测', 2)
add_para('操作步骤：')
add_para('1. 点击顶部"图片检测"Tab')
add_para('2. 点击"选择图片文件夹"，选择施工图纸文件夹')
add_para('3. 系统自动加载并检测第一张图')
add_para('4. 使用左/右箭头或底部胶片条切换图片')
add_para('5. 支持"对比视图"和"仅看结果"两种模式')
add_para('')
add_para('检测结果（右侧面板）：')
add_para('  检测统计：按类别统计目标数量')
add_para('  检测列表：每个目标的类别、置信度')
add_para('  规范检查：12条规则合规/不合规/无法检测 + 总分')
add_para('')
add_para('导出功能：导出图片（带标注框）/ 导出JSON（完整数据）')

add_heading('2.3 文档分析 - 上传', 2)
add_para('1. 点击顶部"文档分析"Tab')
add_para('2. 点击"选择 .docx 文件"上传施工方案文档')
add_para('3. 系统自动提取文档中的图片并展示')
add_para('4. 重复上传同一文件会提示复用历史分析结果')
add_para('')
add_para('历史任务管理：打开（恢复任务）/ 删除（清除所有数据）')

add_heading('2.4 文档分析 - 分析图片', 2)
add_para('单张分析：', bold=True)
add_para('1. 点击左侧图片网格中任意图片')
add_para('2. 在右侧详情面板点击"AI分析此图"')
add_para('3. 等待AI返回分析结果')
add_para('')
add_para('批量分析：', bold=True)
add_para('1. 点击"批量分析全部"或"继续批量分析"')
add_para('2. 系统并发分析所有待处理图片（已成功的自动跳过）')
add_para('3. 实时进度条 + 每张图状态指示')
add_para('')
add_para('队列状态指示：', bold=True)
add_table(
    ['指示', '含义'],
    [
        ['灰点', '等待分析'],
        ['紫色脉冲点 + 旋转遮罩', '正在分析中'],
        ['绿色对勾', '分析完成'],
        ['红色感叹号', '分析异常'],
    ]
)
add_para('"继续批量分析"仅重新处理失败图片，已成功的不重复分析。')

add_heading('2.5 文档分析 - 查看结果', 2)
add_para('选中已分析的图片，右侧详情面板显示：')
add_para('  图纸类型：AI分类结果（彩色标签）')
add_para('  AI摘要：图片内容200-400字详细描述')
add_para('  AI评估：四维度评估（完整性/标注清晰度/布局合理性/规范符合性）+ 优/良/中/差')
add_para('  施工计划：工期和里程碑（如有）')
add_para('  尺寸规格：尺寸、型号、数量（如有）')
add_para('')
add_para('异常处理：', bold=True)
add_para('  网络异常时图片卡片显示红色感叹号，状态栏显示"待分析"数量')
add_para('  可单张重试或使用"继续批量分析"（仅重试失败图）')
add_para('  顶部横幅提示"AI分析未完成（网络连接异常），建议重新分析或人工审核"')

add_heading('2.6 汇总报告与导出', 2)
add_para('批量分析完成后自动生成综合评估报告：')
add_para('  总体概述与整体质量评级（优/良/中/差）')
add_para('  图纸类型分布统计表')
add_para('  质量评级分布（百分比）')
add_para('  各图片评估详情')
add_para('  关键施工参数汇总（桩基规格、分区面积、设备配置、工期等）')
add_para('  AI综合评审意见（可选补充）')
add_para('')
add_para('导出PDF：', bold=True)
add_para('点击汇总报告右上角"导出PDF"按钮 -> 新标签页打开打印版报告 -> 自动弹出打印对话框 -> 浏览器"另存为PDF"保存')
add_para('')
add_para('免责声明：', bold=True)
add_para('本系统在CV的基础上引入大模型用于辅助审查，AI生成可能有误，请仔细甄别。汇总报告的统计数据由系统自动计算，不依赖AI，保证数字准确。')

# ============================================================
# 三、API文档
# ============================================================
add_heading('三、API文档', 1)
add_para('Base URL: http://localhost:5000/api')
add_para('所有接口统一错误格式: {"success": false, "error": "错误描述"}')

add_heading('3.1 文档分析接口（12个端点）', 2)

add_para('POST /docx/upload — 上传.docx，提取图片', bold=True)
add_table(['参数', '类型', '必需', '说明'], [['file', 'File', '是', '.docx文件']])
add_para('响应: {"success":true, "task_id":"...", "total_images":10, "images":[...], "reused":false}')
add_para('reused=true时附带reuse_hint提示复用历史结果。')

add_para('POST /docx/analyze/{task_id} — 启动批量分析', bold=True)
add_para('仅处理未分析或上次异常的图片，已成功的自动跳过。')
add_para('响应: {"success":true, "status":"batch_started", "total":3}')

add_para('POST /docx/analyze-single/{task_id}/{index} — 单张分析', bold=True)
add_para('响应包含完整AI分析结果（image_type/summary/evaluation/elements/schedule/dimensions）。')

add_para('GET /docx/status/{task_id} — 任务状态（HTTP轮询）', bold=True)
add_para('返回batch_running/batch_progress/batch_total/batch_summary/results等字段。')

add_para('GET /docx/status/{task_id}/stream — SSE实时进度（推荐）', bold=True)
add_table(
    ['事件类型', '说明', '附加字段'],
    [
        ['analyzing', '开始分析某张图', 'idx'],
        ['progress', '某张图分析完成', 'idx, result, progress, total, error_count'],
        ['summarizing', '正在生成汇总报告', 'progress, total'],
        ['done', '批量分析完成', 'summary, error_count, status'],
        ['error', '批量分析崩溃', 'error'],
        ['snapshot', '初始快照（断线重连恢复）', 'results, batch_running, batch_summary'],
        ['ping', '心跳保活', '-'],
    ]
)
add_para('SSE连接失败时前端自动降级为HTTP轮询。')

add_para('GET /docx/image/{task_id}/{filename} — 获取提取的图片', bold=True)
add_para('GET /docx/report/{task_id}/html — 打印版HTML报告（自动弹出打印对话框）', bold=True)
add_para('GET /docx/tasks — 列出所有历史任务', bold=True)
add_para('POST /docx/task/{task_id}/load — 加载历史任务到内存', bold=True)
add_para('DELETE /docx/task/{task_id} — 删除任务及所有数据', bold=True)

add_heading('3.2 图片检测接口', 2)

add_para('POST /detection/detect — YOLO目标检测', bold=True)
add_table(
    ['字段', '类型', '说明'],
    [
        ['class', 'string', '类别名称（如"塔吊"）'],
        ['confidence', 'float', '置信度 0-1'],
        ['bbox', 'array[4]', '[x1, y1, x2, y2] 坐标'],
        ['category', 'string', '所属大类（如"垂直运输机械"）'],
    ]
)
add_para('响应还包含class_counts统计、detected_image(base64标注图)、rules_check_results。')

add_para('GET /detection/model/status — 模型加载状态', bold=True)
add_para('POST /detection/switch-model — 切换YOLO模型 (Body: {"model_name":"best0313"})', bold=True)
add_para('GET /api/models/current — 当前模型信息（简化版）', bold=True)
add_para('GET /api/network_check — SiliconFlow API连通性检查', bold=True)

add_heading('3.3 施工规范检查', 2)
add_para('POST /check-rules — 12条规范检查 (Body: {"detections": [...]})', bold=True)

add_table(
    ['ID', '类别', '规则说明', '严重程度'],
    [
        ['R01', '塔吊覆盖范围', '塔吊覆盖范围应覆盖全部施工区域', '严重'],
        ['R02', '塔吊数量', '大型项目至少配置2台塔吊', '重要'],
        ['R03', '施工大门位置', '大门应设置在主要运输通道一侧', '重要'],
        ['R04', '大门-道路连接', '大门必须与临时道路连通', '严重'],
        ['R05', '临时道路完整性', '临时道路应形成环通或尽端回车场', '重要'],
        ['R06', '道路宽度', '主干道>=6m，次要道路>=4m', '一般'],
        ['R07', '办公区布局', '办公区应独立设置，远离施工危险区', '一般'],
        ['R08', '材料堆场位置', '材料堆场应靠近加工区和道路', '重要'],
        ['R09', '钢筋加工区位置', '钢筋加工区应靠近塔吊覆盖范围', '重要'],
        ['R10', '安全设施配置', '应配置沉淀池、洗车槽、门禁系统', '严重'],
        ['R11', '消防设施覆盖', '消防设施应覆盖办公区、生活区、材料区', '严重'],
        ['R12', '临电设施安全距离', '配电设施应远离水源和易燃物', '严重'],
    ]
)

add_heading('3.4 检测类别对照表', 2)
add_table(
    ['英文类别', '中文类别', '所属分类'],
    [
        ['crane', '起重机', '垂直运输机械'],
        ['tower_crane', '塔吊', '垂直运输机械'],
        ['excavator', '挖掘机', '施工机械'],
        ['mixer', '搅拌机', '施工机械'],
        ['rebar_plant', '钢筋加工厂', '临时设施-生产加工区'],
        ['dormitory', '宿舍', '临时设施-生活及办公区'],
        ['office', '办公室', '临时设施-生活及办公区'],
        ['toilet', '厕所', '临时设施-辅助设施'],
        ['gate', '大门', '基础设施'],
        ['red_line', '红线', '基础设施'],
        ['road', '道路', '基础设施'],
        ['stairs', '楼梯', '其他'],
    ]
)

add_heading('3.5 错误响应', 2)
add_table(
    ['状态码', '说明'],
    [
        ['200', '成功'],
        ['400', '请求参数错误（文件格式不支持等）'],
        ['404', '任务或资源不存在'],
        ['413', '文件过大（超过200MB限制）'],
        ['500', '服务器内部错误'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = 'docs/CV检测系统2.0.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
print('Done!')
