"""将三个文档 (SYSTEM_DESIGN.md, USER_MANUAL.md, API.md) 汇总写入一个 docx 文件"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCS_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'docs')
OUTPUT_FILE = os.path.join(OUTPUT_DIR, '项目文档汇总.docx')

FILES = [
    ('系统设计文档', 'SYSTEM_DESIGN.md'),
    ('使用说明书', 'USER_MANUAL.md'),
    ('API 文档', 'API.md'),
]

WESTERN_FONT = 'Times New Roman'
EASTERN_FONT = '微软雅黑'


def set_dual_font(run, western=None, eastern=None, size=None, bold=None, color=None, italic=None):
    """同时设置西文字体和中文字体"""
    w = western or WESTERN_FONT
    e = eastern or EASTERN_FONT
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{w}" w:hAnsi="{w}" w:eastAsia="{e}" w:cs="{w}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), w)
        rFonts.set(qn('w:hAnsi'), w)
        rFonts.set(qn('w:eastAsia'), e)
        rFonts.set(qn('w:cs'), w)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def apply_style_dual_font(style, western=None, eastern=None, size=None):
    """给 Word 样式设置双字体"""
    rpr = style.element.get_or_add_rPr()
    w = western or WESTERN_FONT
    e = eastern or EASTERN_FONT
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{w}" w:hAnsi="{w}" w:eastAsia="{e}" w:cs="{w}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), w)
        rFonts.set(qn('w:hAnsi'), w)
        rFonts.set(qn('w:eastAsia'), e)
        rFonts.set(qn('w:cs'), w)
    if size is not None:
        style.font.size = size


def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def format_table_cell_runs(cell, size=Pt(10), bold=False, align=None):
    """统一格式化表格单元格中的 runs"""
    for p in cell.paragraphs:
        if align:
            p.alignment = align
        for run in p.runs:
            set_dual_font(run, size=size, bold=bold)


def add_table_from_md(doc, lines, start_idx):
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1
    if len(table_lines) < 2:
        return i
    header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    rows = []
    for line in table_lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    if not rows:
        return i
    num_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(header):
        cell = table.rows[0].cells[ci]
        cell.text = h
        format_table_cell_runs(cell, size=Pt(10), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, 'E0E7FF')
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            if ci < num_cols:
                cell = table.rows[ri + 1].cells[ci]
                cell.text = val
                format_table_cell_runs(cell, size=Pt(10))
    doc.add_paragraph()
    return i


def clean_inline(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def add_code_block(doc, code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_dual_font(run, western='Courier New', size=Pt(9), color=RGBColor(0x1E, 0x29, 0x3B))
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        p._element.get_or_add_pPr().append(shading)


def parse_md_to_doc(doc, md_text, title):
    doc.add_heading(title, level=1)
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_buffer = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if stripped == '':
            i += 1
            continue
        if stripped.startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue
        if stripped.startswith('#### '):
            text = clean_inline(stripped[5:])
            p = doc.add_heading(text, level=4)
            i += 1
            continue
        if stripped.startswith('### '):
            text = clean_inline(stripped[4:])
            p = doc.add_heading(text, level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            text = clean_inline(stripped[3:])
            p = doc.add_heading(text, level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            i += 1
            continue
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = clean_inline(stripped[2:])
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            set_dual_font(run, size=Pt(11))
            i += 1
            continue
        match = re.match(r'^(\d+)\.\s+', stripped)
        if match:
            text = clean_inline(stripped[match.end():])
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            set_dual_font(run, size=Pt(11))
            i += 1
            continue
        if stripped.startswith('> '):
            text = clean_inline(stripped[2:])
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            set_dual_font(run, size=Pt(11), italic=True, color=RGBColor(0x47, 0x56, 0x6B))
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
            p._element.get_or_add_pPr().append(shading)
            i += 1
            continue
        if stripped.startswith('---') or stripped.startswith('___') or stripped.startswith('***'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 40)
            set_dual_font(run, size=Pt(8), color=RGBColor(0xCB, 0xD5, 0xE1))
            i += 1
            continue
        text = clean_inline(stripped)
        p = doc.add_paragraph(text)
        for run in p.runs:
            set_dual_font(run, size=Pt(11))
        i += 1


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    apply_style_dual_font(style, size=Pt(11))
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)

    for h in range(1, 5):
        hs = doc.styles[f'Heading {h}']
        apply_style_dual_font(hs)
        hs.font.color.rgb = RGBColor(0x31, 0x2E, 0x8B) if h <= 2 else RGBColor(0x43, 0x3B, 0xB5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run('海南机器管招投标项目')
    set_dual_font(run, size=Pt(26), bold=True, color=RGBColor(0x31, 0x2E, 0x8B))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(6)
    run = sub_p.add_run('建筑图纸分析系统')
    set_dual_font(run, size=Pt(20), bold=True, color=RGBColor(0x63, 0x66, 0xF1))

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(30)
    run = info_p.add_run('项目文档汇总')
    set_dual_font(run, size=Pt(14), color=RGBColor(0x94, 0xA3, 0xB8))

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep.add_run('━' * 50)
    set_dual_font(run, size=Pt(10), color=RGBColor(0xCB, 0xD5, 0xE1))

    toc_p = doc.add_paragraph()
    toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_p.paragraph_format.space_before = Pt(16)
    toc_p.paragraph_format.space_after = Pt(24)
    run = toc_p.add_run('目录')
    set_dual_font(run, size=Pt(14), bold=True)

    toc_items = ['一、系统设计文档', '二、使用说明书', '三、API 文档']
    for item in toc_items:
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tp.add_run(item)
        set_dual_font(run, size=Pt(12), color=RGBColor(0x63, 0x66, 0xF1))

    doc.add_page_break()

    for title_key, filename in FILES:
        filepath = os.path.join(DOCS_DIR, filename)
        if not os.path.exists(filepath):
            print(f'文件不存在: {filepath}')
            continue
        with open(filepath, 'r', encoding='utf-8') as f:
            md_content = f.read()
        parse_md_to_doc(doc, md_content, title_key)
        doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f'文档已生成: {OUTPUT_FILE}')


if __name__ == '__main__':
    main()
