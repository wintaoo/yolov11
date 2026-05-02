from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
doc.styles['Normal'].font.size = Pt(12)

title = doc.add_heading('DZSS项目施工组织设计方案', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

base = r'c:\Users\wint\Desktop\yolov11\demos\demophotos'

images = [
    ('10.1.5施工总平面布置图.jpg', '总平面布置图',
     '图1：施工现场总平面布置图。展示了施工区、办公区、生活区、材料堆场、钢筋加工区、施工大门及临时道路的完整布局，标注清晰，比例准确。'),
    ('10.1.7土方开挖阶段平面布置图.png', '土方工程图',
     '图2：土方开挖阶段平面布置图。基坑开挖深度5.2m，采用1:1放坡+土钉墙支护。标注了开挖范围、出土方向和临时堆土场位置，土方开挖总量约12万立方。'),
    ('12.5.5.1  土方开挖施工阶段平面布置图.png', '施工分区图',
     '图3：土方开挖阶段施工分区图。将场地分为A/B/C三个施工分区，各分区独立流水作业，标注了各分区的边界和作业范围。'),
    ('12.5.5.3  主体结构施工阶段平面布置图.png', '主体结构图',
     '图4：主体结构施工阶段平面布置图。框架-剪力墙体系，柱网8.4m×8.4m。标注了塔吊(TC6015型2台)、施工电梯(SC200/200)、钢筋堆场、PC构件堆场位置。'),
    ('12.5.5.4  装饰装修施工阶段平面布置图.png', '装饰装修图',
     '图5：装饰装修施工阶段平面布置图。含外立面装饰、室内精装修、公共区域装修三个施工区，标注了材料垂直运输通道和周转材料堆放区。'),
    ('12.5.5.5  园林绿化施工阶段平面布置图.png', '周边环境图',
     '图6：园林绿化施工阶段及周边环境图。展示景观施工范围、绿化种植区域、硬质铺装区及与周边道路的衔接关系。'),
    ('12.5.5.6  视频监控平面布置图.png', '临时用电布置图',
     '图7：视频监控及临时用电平面布置图。含监控点位(共24处)、监控室位置、弱电线路敷设路径、配电房及配电箱分布。'),
    ('1.png', '临建设施平面布置图',
     '图8：临建设施平面布置图。含办公区(办公室5间含1间会议室)、生活区(宿舍20间、食堂、卫生间)、门卫室、洗车台、标养室等。'),
    ('8.jpg', '进度计划图',
     '图9：施工进度计划图。标注了各施工阶段的起止时间：施工准备2026.3.1-4.15，土方2026.4.16-6.30，基础2026.7.1-9.30，主体2026.10.1-2027.4.30，装饰装修2027.5.1-11.30。'),
    ('9.jpg', '基础结构图',
     '图10：基础结构平面图。桩基+筏板基础，桩径800mm、桩长28m共216根，筏板厚1.2m C35混凝土。标注了桩位编号、桩顶标高和承台尺寸。'),
]

doc.add_heading('一、工程概况', level=1)
doc.add_paragraph('DZSS项目位于海南地区，总建筑面积约12万平方米，包含8栋高层住宅及配套商业设施。结构形式为框架-剪力墙结构，基础形式为桩基+筏板基础。')
doc.add_paragraph('项目采用EPC总承包模式，由东南大学建筑设计研究院负责设计。合同工期24个月，计划于2026年3月1日开工，2028年2月28日竣工。')

doc.add_heading('二、施工总平面布置', level=1)
doc.add_paragraph('施工总平面布置遵循功能分区明确、物流通道畅通、临建设施集中布置的原则。主要包括施工区、办公区、生活区、材料加工及堆放区四大功能板块。')

doc.add_heading('2.1 总平面布置图', level=2)
img_path, title_text, caption = images[0]
fp = os.path.join(base, img_path)
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(caption)

doc.add_heading('2.2 土方开挖阶段平面布置图', level=2)
_, _, caption = images[1]
fp = os.path.join(base, images[1][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(caption)

doc.add_heading('2.3 施工分区图', level=2)
fp = os.path.join(base, images[2][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[2][2])
doc.add_paragraph('三个分区实行流水施工组织。A区面积约8000、B区约6500、C区约5500。各分区独立配备垂直运输机械。')

doc.add_heading('三、主体结构与装饰装修', level=1)

doc.add_heading('3.1 主体结构施工阶段', level=2)
fp = os.path.join(base, images[3][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[3][2])

doc.add_heading('3.2 装饰装修施工阶段', level=2)
fp = os.path.join(base, images[4][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[4][2])

doc.add_heading('四、园林绿化及周边环境', level=1)
doc.add_paragraph('项目施工前对周边环境进行了详细踏勘，识别出需重点保护的敏感目标，制定了针对性的施工防护措施。')

fp = os.path.join(base, images[5][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[5][2])

doc.add_heading('五、临时设施与附属设施', level=1)

doc.add_heading('5.1 视频监控及临时用电布置', level=2)
fp = os.path.join(base, images[6][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[6][2])

doc.add_heading('5.2 临建设施平面布置', level=2)
fp = os.path.join(base, images[7][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[7][2])

doc.add_heading('六、施工计划与基础工程', level=1)

doc.add_heading('6.1 施工进度计划', level=2)
doc.add_paragraph('施工总工期24个月，分五个阶段组织实施。关键线路为：施工准备-土方工程-基础工程-主体结构-装饰装修。')
fp = os.path.join(base, images[8][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[8][2])

doc.add_heading('6.2 基础结构', level=2)
fp = os.path.join(base, images[9][0])
if os.path.exists(fp):
    doc.add_picture(fp, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(images[9][2])

out = r'c:\Users\wint\Desktop\yolov11\demos\demodocx\施工方案测试文档.docx'
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
sz = os.path.getsize(out)
print(f'Done: {out}')
if sz > 1024*1024:
    print(f'Size: {sz/1024/1024:.1f} MB')
else:
    print(f'Size: {sz/1024:.0f} KB')
for name, _, _ in images:
    fp = os.path.join(base, name)
    ok = 'OK' if os.path.exists(fp) else 'MISS'
    s = os.path.getsize(fp) if os.path.exists(fp) else 0
    if s > 1024*1024:
        si = f'{s/1024/1024:.1f}MB'
    else:
        si = f'{s/1024:.0f}KB'
    print(f'  {ok}  {name}  ({si})')
