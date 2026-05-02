from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.size = Pt(12)

title = doc.add_heading('海南机器管招投标项目——施工方案', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('一、项目概况', level=1)
doc.add_paragraph('本项目为海南地区某大型建筑工程项目，总建筑面积约 12 万平方米，包含 8 栋住宅楼及配套商业设施。计划工期 24 个月，自 2026 年 3 月 1 日至 2028 年 2 月 28 日。')
doc.add_paragraph('项目采用 EPC 总承包模式，由东南大学建筑设计研究院负责设计，海南建工集团负责施工总承包。')

doc.add_heading('二、施工总体部署', level=1)
doc.add_paragraph('施工总体按照先地下后地上、先主体后装修、先土建后安装的原则进行组织。施工场地划分为三个施工分区：A区（1-3#楼）、B区（4-6#楼）、C区（7-8#楼及商业配套）。')

doc.add_heading('2.1 施工分区规划图', level=2)
doc.add_paragraph('下图为施工分区规划图，展示了三个施工分区的范围和位置关系：')

img_dir = r'c:\Users\wint\Desktop\yolov11\runs\detect\yolov11l_custom8'
seg_dir = r'c:\Users\wint\Desktop\yolov11\runs\segment\yolov11l_custom8'

img1 = os.path.join(img_dir, 'train_batch0.jpg')
if os.path.exists(img1):
    doc.add_picture(img1, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图1：施工分区规划图。图中展示A区（左侧）、B区（中部）、C区（右侧）三个分区。A区占地约8000㎡，B区占地约6500㎡，C区占地约5500㎡。')

doc.add_heading('2.2 施工进度计划图', level=2)
doc.add_paragraph('施工总体进度计划采用横道图形式，关键节点如下：')

img2 = os.path.join(img_dir, 'train_batch1.jpg')
if os.path.exists(img2):
    doc.add_picture(img2, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图2：施工进度计划横道图。主体结构施工计划于 2026年6月1日至 2027年3月31日，装饰装修阶段计划于 2027年4月1日至 2027年11月30日。')

doc.add_heading('2.3 总平面布置图', level=2)
doc.add_paragraph('施工现场总平面布置图包含施工道路、材料堆场、加工区、办公生活区等临时设施：')

img3 = os.path.join(img_dir, 'train_batch2.jpg')
if os.path.exists(img3):
    doc.add_picture(img3, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图3：施工现场总平面布置图。图中标注了施工大门（1#大门、2#大门）、围挡、办公区（办公室5间、会议室1间）、生活区（宿舍20间、食堂、卫生间）、钢筋加工区、木工加工区、材料堆放区及施工道路。')

doc.add_heading('三、临建设施布置', level=1)

doc.add_heading('3.1 临时用电布置图', level=2)
doc.add_paragraph('施工现场临时用电从业主提供的 630KVA 箱变接入，设置总配电房1座，分配电箱6处。')

img4 = os.path.join(img_dir, 'train_batch5040.jpg')
if os.path.exists(img4):
    doc.add_picture(img4, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图4：临时用电平面布置图。1#配电房位于场地西北角，2#配电房位于场地东南角，临电线路沿施工道路一侧敷设。')

doc.add_heading('3.2 临时用水布置图', level=2)
doc.add_paragraph('施工用水从市政DN200给水管接入，场内设置供水立管2处、消防栓8处。')

img5 = os.path.join(img_dir, 'train_batch5041.jpg')
if os.path.exists(img5):
    doc.add_picture(img5, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图5：临时用水平面布置图。供水主管沿施工道路敷设，消防栓均匀布置在各施工分区。')

doc.add_heading('四、土方及基础工程', level=1)
doc.add_paragraph('本工程基坑开挖深度约 5.2m，采用放坡+土钉墙支护方案。土方开挖总量约 12 万 m³。')

img6 = os.path.join(img_dir, 'train_batch5042.jpg')
if os.path.exists(img6):
    doc.add_picture(img6, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图6：土方工程图/基础结构图。展示基坑开挖范围及支护结构布置。基础采用桩基+筏板基础。')

doc.add_heading('五、周边环境图', level=1)
doc.add_paragraph('项目地块北侧为城市主干道，东侧为已建成住宅小区，南侧为规划市政公园，西侧为空地。')

img7 = os.path.join(seg_dir, 'train_batch0.jpg')
if os.path.exists(img7):
    doc.add_picture(img7, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图7：项目区位图/周边环境图。展示项目位置与周边道路、建筑的关系。')

doc.add_heading('六、各阶段施工平面布置', level=1)

doc.add_heading('6.1 施工准备阶段', level=2)
img8 = os.path.join(seg_dir, 'train_batch1.jpg')
if os.path.exists(img8):
    doc.add_picture(img8, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图8：施工准备阶段平面布置图。完成场地平整、临时道路、围挡及临建设施搭建。')

doc.add_heading('6.2 土方阶段', level=2)
img9 = os.path.join(seg_dir, 'train_batch2.jpg')
if os.path.exists(img9):
    doc.add_picture(img9, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图9：土方阶段平面布置图。布置土方开挖机械、运输通道及临时堆土场。')

doc.add_heading('6.3 主体结构阶段', level=2)
img10 = os.path.join(img_dir, 'val_batch0_labels.jpg')
if os.path.exists(img10):
    doc.add_picture(img10, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('图10：主体结构阶段平面布置图。布置塔吊（TC6015型号2台）、施工电梯（SC200/200型号）、钢筋堆场、PC构件堆场及材料周转区。')

output_path = r'c:\Users\wint\Desktop\yolov11\demos\demodocx\测试文档.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'测试文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
